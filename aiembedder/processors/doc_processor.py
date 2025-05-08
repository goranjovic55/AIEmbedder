"""
DOC/DOCX document processor for AIEmbedder.
"""

from typing import Optional

from docx import Document

from aiembedder.processors.base_processor import BaseProcessor
from aiembedder.utils.errors import ProcessingError
from aiembedder.utils.logging import Logger

class DocProcessor(BaseProcessor):
    """Processor for DOC/DOCX documents."""
    
    def __init__(self, logger: Optional[Logger] = None):
        """Initialize DOC processor.
        
        Args:
            logger: Logger instance
        """
        super().__init__(logger)
        self.logger.info("Initialized DOC processor")
    
    def process(self, file_path: str) -> str:
        """Process DOC/DOCX document and return text content.
        
        Args:
            file_path: Path to DOC/DOCX file
            
        Returns:
            Extracted text content
            
        Raises:
            FileAccessError: If file cannot be accessed
            ProcessingError: If processing fails
        """
        try:
            self.validate_file(file_path)
            self.logger.info(f"Processing DOC file: {file_path}")
            
            # Load document
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            
            # Join paragraphs with newlines
            text = '\n'.join(paragraphs)
            
            if not text:
                raise ProcessingError("No text content extracted", "PROC_001")
            
            self.logger.info(f"Successfully processed DOC file: {file_path}")
            return text
            
        except ProcessingError:
            raise
        except Exception as e:
            self.log_error(e, file_path)
            raise ProcessingError(f"Failed to process DOC file: {str(e)}", "PROC_001") 